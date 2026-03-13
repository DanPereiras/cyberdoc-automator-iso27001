import os
import re
import warnings
from datetime import datetime
from tqdm import tqdm
from openpyxl import load_workbook, Workbook
from openpyxl.styles import Font, Alignment, PatternFill
from docx import Document
from docx.shared import Mm

# Ocultar avisos para manter o terminal limpo
warnings.simplefilter('ignore', UserWarning)

# Configurações Globais
REGEX_BRACKETS = r'\[(.*?)\]'

class CyberDocAutomator:
    def __init__(self, pasta_raiz):
        self.pasta_raiz = os.path.normpath(pasta_raiz)
        self.logs = []

    def listar_arquivos(self, extensoes=('.docx',)):
        arquivos = []
        for root, _, filenames in os.walk(self.pasta_raiz):
            for f in filenames:
                if not f.startswith('~$') and f.lower().endswith(extensoes):
                    arquivos.append(os.path.join(root, f))
        return arquivos

    # --- MÓDULO DE EXTRAÇÃO ---
    def extrair_termos(self):
        print(f"\n🔍 Iniciando extração de termos em: {self.pasta_raiz}")
        arquivos = self.listar_arquivos(('.docx', '.xlsx'))
        resultados = []

        for caminho in tqdm(arquivos, desc="Analisando"):
            try:
                termos = []
                if caminho.endswith('.docx'):
                    doc = Document(caminho)
                    termos.extend(self._buscar_em_doc(doc))
                elif caminho.endswith('.xlsx'):
                    wb = load_workbook(caminho, data_only=True)
                    for sheet in wb.sheetnames:
                        for row in wb[sheet].iter_rows():
                            for cell in row:
                                val = str(cell.value) if cell.value else ""
                                termos.extend(re.findall(REGEX_BRACKETS, val))
                
                if termos:
                    resultados.append({'arquivo': os.path.basename(caminho), 'termos': termos})
            except Exception as e:
                print(f"⚠️ Erro em {os.path.basename(caminho)}: {e}")

        return self._salvar_log_extracao(resultados)

    def _buscar_em_doc(self, doc):
        termos = []
        # Parágrafos, Tabelas, Cabeçalhos e Rodapés
        for p in doc.paragraphs: termos.extend(re.findall(REGEX_BRACKETS, p.text))
        for t in doc.tables:
            for r in t.rows:
                for c in r.cells:
                    for p in c.paragraphs: termos.extend(re.findall(REGEX_BRACKETS, p.text))
        for s in doc.sections:
            for hf in [s.header, s.footer]:
                for p in hf.paragraphs: termos.extend(re.findall(REGEX_BRACKETS, p.text))
        return termos

    def _salvar_log_extracao(self, resultados):
        path = os.path.join(self.pasta_raiz, f"mapeamento_termos_{datetime.now().strftime('%Y%m%d_%H%M%S')}.xlsx")
        wb = Workbook()
        ws = wb.active
        ws.title = "Termos Encontrados"
        ws.append(["Arquivo", "Termo [Tags]", "Ocorrências"])
        
        for res in resultados:
            contagem = {}
            for t in res['termos']: contagem[t] = contagem.get(t, 0) + 1
            for termo, qtd in contagem.items():
                ws.append([res['arquivo'], f"[{termo}]", qtd])
        
        wb.save(path)
        return path

    # --- MÓDULO DE SUBSTITUIÇÃO ---
    def processar_substituicao(self, config):
        arquivos = self.listar_arquivos(('.docx',))
        self.logs = []

        for caminho in tqdm(arquivos, desc="Processando"):
            try:
                doc = Document(caminho)
                alterado = False
                t_txt, t_img, hist = 0, 0, False

                if config.get('texto'):
                    t_txt = self._substituir_texto(doc, config['alvo'], config['novo'])
                
                if config.get('imagem'):
                    t_img = self._substituir_imagem(doc, config['alvo_img'], config['path_img'], config['width'])
                
                if config.get('versao'):
                    hist = self._atualizar_historico(doc, config['v_num'], config['v_autor'], config['v_desc'])

                if t_txt > 0 or t_img > 0 or hist:
                    doc.save(caminho)
                    self.logs.append({'arquivo': os.path.basename(caminho), 'txt': t_txt, 'img': t_img, 'hist': hist})
            except Exception as e:
                print(f"⚠️ Erro ao processar {os.path.basename(caminho)}: {e}")

        return self._salvar_log_processamento()

    def _substituir_texto(self, doc, alvo, novo):
        total = 0
        tag = f"[{alvo}]"
        
        def replace_in_paragraphs(paragraphs):
            count = 0
            for p in paragraphs:
                if tag in p.text:
                    count += p.text.count(tag)
                    # Preserva formatação básica ao substituir no primeiro run
                    p.text = p.text.replace(tag, novo)
            return count

        total += replace_in_paragraphs(doc.paragraphs)
        for t in doc.tables:
            for r in t.rows:
                for c in r.cells: total += replace_in_paragraphs(c.paragraphs)
        for s in doc.sections:
            for hf in [s.header, s.footer]: total += replace_in_paragraphs(hf.paragraphs)
        return total

    def _substituir_imagem(self, doc, alvo, path_img, width):
        total = 0
        tag = f"[{alvo}]"
        
        def replace_with_img(paragraphs):
            count = 0
            for p in paragraphs:
                if tag in p.text:
                    count += p.text.count(tag)
                    p.text = p.text.replace(tag, "")
                    run = p.add_run()
                    run.add_picture(path_img, width=Mm(width))
            return count

        total += replace_with_img(doc.paragraphs)
        for t in doc.tables:
            for r in t.rows:
                for c in r.cells: total += replace_with_img(c.paragraphs)
        return total

    def _atualizar_historico(self, doc, v, autor, desc):
        for t in doc.tables:
            header = [c.text.lower() for c in t.rows[0].cells]
            if any("versão" in x for x in header) and any("descrição" in x for x in header):
                row = t.add_row().cells
                row[0].text = datetime.now().strftime("%d/%m/%Y")
                row[1].text = v
                row[2].text = autor
                row[3].text = desc
                return True
        return False

    def _salvar_log_processamento(self):
        path = os.path.join(self.pasta_raiz, f"log_execucao_{datetime.now().strftime('%Y%m%d_%H%M%S')}.xlsx")
        wb = Workbook()
        ws = wb.active
        ws.append(["Arquivo", "Texto", "Imagens", "Histórico"])
        for l in self.logs:
            ws.append([l['arquivo'], l['txt'], l['img'], "Sim" if l['hist'] else "Não"])
        wb.save(path)
        return path

def menu():
    print("\n" + "="*50)
    print("🛡️  CYBERDOC AUTOMATOR PRO - ISO 27001 EDITION")
    print("="*50)
    pasta = input("📁 Caminho da pasta de documentos: ").strip().replace('"', '')
    if not os.path.exists(pasta):
        print("❌ Pasta não encontrada!"); return

    app = CyberDocAutomator(pasta)
    
    while True:
        print("\n[1] Mapear Termos [Tags] existentes")
        print("[2] Substituir Texto em Massa")
        print("[3] Substituir por Imagem (Logo/Assinatura)")
        print("[4] Combo: Texto + Imagem + Nova Versão")
        print("[s] Sair")
        
        op = input("\nEscolha: ").lower()
        if op == 's': break
        
        if op == '1':
            log = app.extrair_termos()
            print(f"✅ Mapeamento concluído: {log}")
        
        elif op in ['2', '3', '4']:
            conf = {'texto': op in ['2', '4'], 'imagem': op in ['3', '4'], 'versao': op == '4'}
            if conf['texto']:
                conf['alvo'] = input("🔍 Tag alvo (sem colchetes): ")
                conf['novo'] = input("✏️  Novo texto: ")
            if conf['imagem']:
                conf['alvo_img'] = input("🖼️  Tag para imagem: ")
                conf['path_img'] = input("📂 Caminho da imagem: ").strip().replace('"', '')
                w = input("📏 Largura (mm) [40]: ")
                conf['width'] = int(w) if w.isdigit() else 40
            if conf['versao']:
                conf['v_num'] = input("🔢 Versão: ")
                conf['v_autor'] = input("👤 Autor: ")
                conf['v_desc'] = input("📄 Descrição: ")
            
            log = app.processar_substituicao(conf)
            print(f"✅ Processamento concluído! Log: {log}")

if __name__ == "__main__":
    menu()
